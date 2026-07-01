"""
SOA Word Document Generator
===========================

Builds a Statement of Advice (SOA) skeleton as a Microsoft Word (.docx) file
with EVERY strategy scoped in. The document is structured purely from the SOA's
table-of-contents headings and sub-headings, using native Word Heading styles so
that Word's Navigation Pane (View > Navigation Pane) is fully populated and
clickable.

Under every heading a "Notes" placeholder paragraph is inserted so the adviser
can type notes for each section of the SOA.

Usage:
    python generate_soa_word.py [output_path]

Requires: python-docx  (pip install python-docx)
"""

import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: F401 (kept for extensions)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


# ---------------------------------------------------------------------------
# SOA structure — the full table of contents with ALL strategies scoped in.
#
# Each item is (level, title):
#   level 1 -> Heading 1 (major section)
#   level 2 -> Heading 2 (sub-section / individual strategy)
#   level 3 -> Heading 3 (detail within a strategy)
# ---------------------------------------------------------------------------

SOA_STRUCTURE = [
    (1, "Important Information"),
    (2, "Purpose of this Statement of Advice"),
    (2, "Your Adviser and Licensee Details"),
    (2, "Basis on Which Advice is Provided"),

    (1, "Executive Summary"),
    (2, "Summary of Your Goals"),
    (2, "Summary of Our Recommendations"),
    (2, "Summary of Benefits and Costs"),

    (1, "Scope of Advice"),
    (2, "What We Have Agreed to Advise On (In Scope)"),
    (2, "What Is Not Included (Out of Scope)"),
    (2, "Information You Have Provided"),

    (1, "About You — Your Current Situation"),
    (2, "Personal and Family Details"),
    (2, "Your Goals and Objectives"),
    (2, "Income and Expenses (Cashflow)"),
    (2, "Assets and Liabilities"),
    (2, "Superannuation Summary"),
    (2, "Existing Insurance Cover"),
    (2, "Estate Planning Arrangements"),
    (2, "Centrelink and Government Benefits"),
    (2, "Your Risk Profile and Attitude to Risk"),

    (1, "Our Recommended Strategies"),
    (2, "Strategy 1 — Cashflow and Budgeting"),
    (2, "Strategy 2 — Debt Management and Debt Recycling"),
    (2, "Strategy 3 — Emergency Fund and Cash Reserve"),
    (2, "Strategy 4 — Superannuation Consolidation"),
    (2, "Strategy 5 — Concessional (Pre-Tax) Super Contributions"),
    (3, "Salary Sacrifice Contributions"),
    (3, "Personal Deductible Contributions"),
    (3, "Carry-Forward (Catch-Up) Concessional Contributions"),
    (2, "Strategy 6 — Non-Concessional (After-Tax) Super Contributions"),
    (3, "Bring-Forward Rule"),
    (3, "Spouse Contributions"),
    (3, "Government Co-Contribution"),
    (2, "Strategy 7 — Downsizer Contributions"),
    (2, "Strategy 8 — Transition to Retirement (TTR) Pension"),
    (2, "Strategy 9 — Account-Based Pension / Retirement Income Stream"),
    (2, "Strategy 10 — Investment Strategy and Portfolio Construction"),
    (2, "Strategy 11 — Capital Gains Tax (CGT) Management"),
    (2, "Strategy 12 — Tax Planning and Optimisation"),
    (2, "Strategy 13 — Personal Insurance"),
    (3, "Life (Death) Cover"),
    (3, "Total and Permanent Disability (TPD) Cover"),
    (3, "Income Protection Cover"),
    (3, "Trauma / Critical Illness Cover"),
    (2, "Strategy 14 — Centrelink / Age Pension Optimisation"),
    (2, "Strategy 15 — Aged Care Planning"),
    (2, "Strategy 16 — Estate Planning"),

    (1, "Product Recommendations"),
    (2, "Recommended Superannuation Fund / Platform"),
    (2, "Recommended Investment Options"),
    (2, "Recommended Insurance Products"),
    (2, "Product Replacement Comparison"),

    (1, "Costs, Fees and Remuneration"),
    (2, "Advice Preparation Fees"),
    (2, "Implementation Fees"),
    (2, "Ongoing Service Fees"),
    (2, "Product and Investment Costs"),
    (2, "How We Are Paid"),

    (1, "Risks, Consequences and Considerations"),
    (2, "General Risks"),
    (2, "Strategy-Specific Risks"),
    (2, "Consequences of Replacing Existing Products"),
    (2, "Things You Should Consider"),

    (1, "Alternative Strategies Considered"),

    (1, "Assumptions"),

    (1, "Implementation Plan and Next Steps"),
    (2, "Actions We Will Complete for You"),
    (2, "Actions You Need to Complete"),
    (2, "Implementation Timeline"),

    (1, "Ongoing Service and Review"),

    (1, "Authority to Proceed"),

    (1, "Appendices and Disclaimers"),
    (2, "Financial Services Guide"),
    (2, "Product Disclosure Statements"),
    (2, "Glossary of Terms"),
]

NOTE_PROMPT = "Notes: "


def add_toc_field(document):
    """Insert a real Word Table of Contents field (updates on open / F9)."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    # Include heading levels 1-3 in the TOC, with hyperlinks and page numbers.
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose “Update Field” to build the table of contents."

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)


def add_note_placeholder(document):
    """Add an italic, grey 'Notes:' line the adviser can type into."""
    para = document.add_paragraph()
    run = para.add_run(NOTE_PROMPT)
    run.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(12)
    return para


def build_document(output_path):
    document = Document()

    # ---- Cover / title page -------------------------------------------------
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Statement of Advice")
    run.bold = True
    run.font.size = Pt(28)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("MyWealth — Prepared for [Client Name]")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Adviser: [Adviser Name]   |   Date: [DD Month YYYY]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    document.add_page_break()

    # ---- Table of contents --------------------------------------------------
    toc_heading = document.add_heading("Table of Contents", level=1)
    # Keep the TOC heading itself out of the navigation clutter? No — leave it in.
    add_toc_field(document)
    document.add_page_break()

    # ---- Body: all headings + note placeholders -----------------------------
    for level, heading_text in SOA_STRUCTURE:
        document.add_heading(heading_text, level=level)
        add_note_placeholder(document)

    document.save(output_path)
    return output_path


if __name__ == "__main__":
    out = sys.argv[1] if len(sys.argv) > 1 else "MyWealth_SOA_Navigation.docx"
    path = build_document(out)
    section_count = sum(1 for lvl, _ in SOA_STRUCTURE if lvl == 1)
    print(f"Generated: {path}")
    print(f"Total headings: {len(SOA_STRUCTURE)} ({section_count} top-level sections)")
